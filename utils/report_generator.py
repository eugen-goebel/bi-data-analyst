"""
Report Generator — Produces a professional DOCX business intelligence report.

Transforms data summaries, pattern analysis, charts, and AI insights into
a formatted Word document with embedded matplotlib visualizations.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.data_loader import DataSummary
from agents.pattern_agent import PatternAnalysis
from agents.visualization_agent import VisualizationResult
from agents.insight_agent import InsightResult


# ---------------------------------------------------------------------------
# Color palette (corporate orange/teal theme)
# ---------------------------------------------------------------------------
COLOR_DARK_ORANGE  = RGBColor(0xC4, 0x49, 0x00)   # #C44900 — headers
COLOR_TEAL         = RGBColor(0x0D, 0x73, 0x77)   # #0D7377 — accents
COLOR_LIGHT_BG     = RGBColor(0xFF, 0xF8, 0xF0)   # #FFF8F0 — table backgrounds
COLOR_WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_PRIORITY_HIGH = "C62828"
COLOR_PRIORITY_MED  = "E65100"
COLOR_PRIORITY_LOW  = "2E7D32"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str):
    """Set a table cell's background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled section heading."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.font.color.rgb = COLOR_DARK_ORANGE
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_TEAL
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    return para


def _add_bullet_list(doc: Document, items: list[str], indent: int = 0):
    """Add a bulleted list of items."""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        run = para.add_run(item)
        run.font.size = Pt(10.5)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal divider."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C44900")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)


def _add_chart(doc: Document, chart_dir: str, chart_config):
    """Embed a chart image with caption."""
    chart_path = os.path.join(chart_dir, chart_config.filename)
    if not os.path.exists(chart_path):
        return

    doc.add_picture(chart_path, width=Inches(5.8))

    # Caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(chart_config.description)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    caption.paragraph_format.space_after = Pt(12)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover_page(doc: Document, summary: DataSummary):
    """Create a styled cover page."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BUSINESS INTELLIGENCE REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK_ORANGE

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run(summary.filename)
    run2.font.size = Pt(20)
    run2.font.bold = True
    run2.font.color.rgb = COLOR_TEAL

    doc.add_paragraph()

    # Stats line
    stats = doc.add_paragraph()
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = stats.add_run(
        f"{summary.row_count:,} records  •  {summary.column_count} columns  •  "
        f"Quality: {summary.data_quality_score}/100"
    )
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run4.font.italic = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = sub.add_run("AI-Powered Data Analysis  •  Multi-Agent BI Pipeline")
    run5.font.size = Pt(10)
    run5.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run5.font.italic = True

    doc.add_page_break()


def _build_metrics_table(doc: Document, stats):
    """Build a table of key numeric statistics."""
    headers = ["Metric", "Mean", "Median", "Std Dev", "Min", "Max"]
    table = doc.add_table(rows=1 + len(stats), cols=6)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.3), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, "C44900")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    # Data rows
    for row_idx, stat in enumerate(stats):
        row = table.rows[row_idx + 1]
        bg = "FFF8F0" if row_idx % 2 == 0 else "FFFFFF"

        values = [
            stat.column.replace("_", " ").title(),
            f"{stat.mean:,.1f}",
            f"{stat.median:,.1f}",
            f"{stat.std:,.1f}",
            f"{stat.min:,.1f}",
            f"{stat.max:,.1f}",
        ]
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            _set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True


def _build_recommendations_table(doc: Document, recommendations):
    """Build a color-coded recommendations table."""
    headers = ["Priority", "Recommendation", "Category", "Expected Impact"]
    table = doc.add_table(rows=1 + len(recommendations), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(0.7), Inches(2.8), Inches(1.2), Inches(1.8)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, "C44900")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    priority_colors = {"High": COLOR_PRIORITY_HIGH, "Medium": COLOR_PRIORITY_MED, "Low": COLOR_PRIORITY_LOW}

    for row_idx, rec in enumerate(recommendations):
        row = table.rows[row_idx + 1]
        bg = "FFF8F0" if row_idx % 2 == 0 else "FFFFFF"

        # Priority cell (color-coded)
        priority_cell = row.cells[0]
        _set_cell_background(priority_cell, priority_colors.get(rec.priority, "FFFFFF"))
        para = priority_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(rec.priority)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_WHITE

        # Title + description
        rec_cell = row.cells[1]
        _set_cell_background(rec_cell, bg)
        para = rec_cell.paragraphs[0]
        run_title = para.add_run(rec.title + "\n")
        run_title.font.bold = True
        run_title.font.size = Pt(9.5)
        run_desc = para.add_run(rec.description)
        run_desc.font.size = Pt(8.5)
        run_desc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Category
        cat_cell = row.cells[2]
        _set_cell_background(cat_cell, bg)
        para = cat_cell.paragraphs[0]
        run = para.add_run(rec.category)
        run.font.size = Pt(9)

        # Impact
        impact_cell = row.cells[3]
        _set_cell_background(impact_cell, bg)
        para = impact_cell.paragraphs[0]
        run = para.add_run(rec.expected_impact)
        run.font.size = Pt(9)
        run.font.italic = True


def _build_outlier_table(doc: Document, outliers):
    """Build a table of detected outliers."""
    if not outliers:
        doc.add_paragraph("No significant outliers detected.")
        return

    headers = ["Column", "Value", "Z-Score", "Description"]
    table = doc.add_table(rows=1 + len(outliers), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.2), Inches(1.0), Inches(0.8), Inches(3.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_background(cell, "0D7377")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    for row_idx, outlier in enumerate(outliers):
        row = table.rows[row_idx + 1]
        bg = "FFF8F0" if row_idx % 2 == 0 else "FFFFFF"

        values = [
            outlier.column.replace("_", " ").title(),
            f"{outlier.value:,.1f}",
            f"{outlier.z_score:.1f}",
            outlier.description,
        ]
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            _set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx_report(
    summary: DataSummary,
    patterns: PatternAnalysis,
    viz_result: VisualizationResult,
    insights: InsightResult,
    output_dir: str = "output",
) -> str:
    """
    Generate a professional DOCX business intelligence report.

    Args:
        summary:    DataSummary from DataLoaderAgent
        patterns:   PatternAnalysis from PatternAgent
        viz_result: VisualizationResult with chart paths
        insights:   InsightResult from InsightAgent (or mock)
        output_dir: Directory where the file will be saved

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    # 1. Cover Page
    _build_cover_page(doc, summary)

    # 2. Executive Summary
    _add_heading(doc, "Executive Summary")
    _add_horizontal_rule(doc)
    para = doc.add_paragraph(insights.executive_summary)
    para.paragraph_format.space_after = Pt(12)

    # 3. Data Overview
    _add_heading(doc, "Data Overview")
    _add_horizontal_rule(doc)
    overview_lines = [
        f"File: {summary.filename}",
        f"Records: {summary.row_count:,} rows × {summary.column_count} columns",
        f"Date Range: {summary.date_range or 'N/A'}",
        f"Data Quality Score: {summary.data_quality_score}/100",
    ]
    for line in overview_lines:
        doc.add_paragraph(line)

    # 4. Key Metrics Summary
    _add_heading(doc, "Key Metrics Summary")
    _add_horizontal_rule(doc)
    _build_metrics_table(doc, summary.numeric_stats)
    doc.add_paragraph()

    # 5. Trend Analysis + Chart
    _add_heading(doc, "Trend Analysis")
    _add_horizontal_rule(doc)
    for trend in patterns.trends[:5]:
        doc.add_paragraph(f"• {trend.description}")
    doc.add_paragraph()

    # Embed line chart
    line_charts = [c for c in viz_result.charts if c.chart_type == "line"]
    for chart in line_charts:
        _add_chart(doc, viz_result.chart_dir, chart)

    # 6. Category & Regional Breakdown + Charts
    _add_heading(doc, "Category & Regional Breakdown")
    _add_horizontal_rule(doc)

    bar_charts = [c for c in viz_result.charts if c.chart_type == "bar"]
    for chart in bar_charts:
        _add_chart(doc, viz_result.chart_dir, chart)

    pie_charts = [c for c in viz_result.charts if c.chart_type == "pie"]
    for chart in pie_charts:
        _add_chart(doc, viz_result.chart_dir, chart)

    # 7. Correlation Analysis + Heatmap
    _add_heading(doc, "Correlation Analysis")
    _add_horizontal_rule(doc)
    for corr in patterns.correlations[:5]:
        doc.add_paragraph(f"• {corr.interpretation}")
    doc.add_paragraph()

    heatmap_charts = [c for c in viz_result.charts if c.chart_type == "heatmap"]
    for chart in heatmap_charts:
        _add_chart(doc, viz_result.chart_dir, chart)

    # 8. Outlier Detection
    _add_heading(doc, "Outlier Detection")
    _add_horizontal_rule(doc)
    _build_outlier_table(doc, patterns.outliers[:10])
    doc.add_paragraph()

    # Scatter chart
    scatter_charts = [c for c in viz_result.charts if c.chart_type == "scatter"]
    for chart in scatter_charts:
        _add_chart(doc, viz_result.chart_dir, chart)

    # 9. Key Findings
    _add_heading(doc, "Key Findings")
    _add_horizontal_rule(doc)
    for finding in insights.key_findings:
        _add_heading(doc, finding.finding, level=2)
        doc.add_paragraph(f"Evidence: {finding.evidence}")
        para = doc.add_paragraph(f"Implication: {finding.business_implication}")
        para.paragraph_format.space_after = Pt(10)

    # 10. Recommendations
    _add_heading(doc, "Recommendations")
    _add_horizontal_rule(doc)
    _build_recommendations_table(doc, insights.recommendations)
    doc.add_paragraph()

    # 11. Risk Alerts
    _add_heading(doc, "Risk Alerts")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, insights.risk_alerts)

    # 12. Opportunities
    _add_heading(doc, "Opportunities")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, insights.opportunities)

    # 13. Methodology
    _add_heading(doc, "Methodology")
    _add_horizontal_rule(doc)
    doc.add_paragraph(insights.methodology_note)

    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated by BI Data Analyst Agent  •  {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # Save
    safe_name = summary.filename.replace(".csv", "").replace(".xlsx", "")
    safe_name = safe_name.lower().replace(" ", "_").replace("/", "-")
    filename = f"bi_report_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
