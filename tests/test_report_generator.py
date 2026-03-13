"""
Tests for the report generator.

Generates a full DOCX report using real sample data plus mock insights,
then validates the output file structure and contents.
"""

import os
import pytest
from docx import Document

from agents.data_loader import DataLoaderAgent
from agents.pattern_agent import PatternAgent
from agents.visualization_agent import VisualizationAgent
from agents.mock_data import MOCK_INSIGHTS
from utils.report_generator import generate_docx_report

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(PROJECT_DIR, "data")
SAMPLE_CSV = os.path.join(DATA_DIR, "sample_sales.csv")


@pytest.fixture(scope="module")
def report_output(tmp_path_factory):
    """Generate a complete report once for the module."""
    loader = DataLoaderAgent()
    summary, df = loader.load(SAMPLE_CSV)

    pattern_agent = PatternAgent()
    patterns = pattern_agent.analyze(df, summary)

    chart_dir = str(tmp_path_factory.mktemp("charts"))
    viz_agent = VisualizationAgent()
    viz_result = viz_agent.create_charts(df, summary, patterns, chart_dir)

    output_dir = str(tmp_path_factory.mktemp("report_output"))
    report_path = generate_docx_report(
        summary=summary,
        patterns=patterns,
        viz_result=viz_result,
        insights=MOCK_INSIGHTS,
        output_dir=output_dir,
    )
    return report_path, output_dir


class TestReportGenerator:
    def test_docx_file_exists(self, report_output):
        report_path, _ = report_output
        assert os.path.exists(report_path)

    def test_docx_file_not_empty(self, report_output):
        report_path, _ = report_output
        assert os.path.getsize(report_path) > 0

    def test_docx_extension(self, report_output):
        report_path, _ = report_output
        assert report_path.endswith(".docx")

    def test_filename_contains_sample_sales(self, report_output):
        report_path, _ = report_output
        basename = os.path.basename(report_path)
        assert "sample_sales" in basename

    def test_report_contains_executive_summary_section(self, report_output):
        report_path, _ = report_output
        doc = Document(report_path)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Executive Summary" in h for h in headings)

    def test_report_contains_key_sections(self, report_output):
        report_path, _ = report_output
        doc = Document(report_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        expected_sections = [
            "Executive Summary",
            "Data Overview",
            "Key Metrics",
            "Trend Analysis",
            "Correlation Analysis",
            "Outlier Detection",
            "Key Findings",
            "Recommendations",
            "Risk Alerts",
            "Opportunities",
            "Methodology",
        ]
        for section in expected_sections:
            assert section in all_text, f"Missing section: {section}"

    def test_report_contains_tables(self, report_output):
        report_path, _ = report_output
        doc = Document(report_path)
        assert len(doc.tables) >= 2, "Report should contain at least metrics and recommendations tables"

    def test_report_contains_images(self, report_output):
        """Check that charts are embedded as images."""
        report_path, _ = report_output
        doc = Document(report_path)
        # Images are stored as inline shapes in python-docx
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        assert image_count >= 3, f"Expected at least 3 embedded images, found {image_count}"

    def test_returns_absolute_path(self, report_output):
        report_path, _ = report_output
        assert os.path.isabs(report_path)

    def test_creates_output_directory(self, tmp_path):
        """generate_docx_report should create output_dir if missing."""
        loader = DataLoaderAgent()
        summary, df = loader.load(SAMPLE_CSV)
        pattern_agent = PatternAgent()
        patterns = pattern_agent.analyze(df, summary)

        chart_dir = str(tmp_path / "charts")
        viz_agent = VisualizationAgent()
        viz_result = viz_agent.create_charts(df, summary, patterns, chart_dir)

        new_output = str(tmp_path / "brand_new_dir")
        assert not os.path.exists(new_output)

        report_path = generate_docx_report(
            summary=summary,
            patterns=patterns,
            viz_result=viz_result,
            insights=MOCK_INSIGHTS,
            output_dir=new_output,
        )
        assert os.path.isdir(new_output)
        assert os.path.exists(report_path)

    def test_report_contains_data_quality_score(self, report_output):
        report_path, _ = report_output
        doc = Document(report_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Quality" in all_text or "quality" in all_text

    def test_report_has_cover_page_title(self, report_output):
        report_path, _ = report_output
        doc = Document(report_path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "BUSINESS INTELLIGENCE REPORT" in all_text
